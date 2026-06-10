"""T5 —— 巡检任务 Word 报告生成。

入口:``build_task_report_docx(...) -> bytes``

该函数是**纯函数**:只接收已组装好的 dict / list,不碰数据库、不碰 ORM,
因此可在 tests/test_reports.py 里用假数据直接测,无需起 DB。

调用方(backend/app/api/reports.py)负责:
  1. 鉴权 + ownership(``Depends(get_owned_task)``)
  2. 从 DB 把 task / results / nests 查出来,转成下面的 dict 结构
  3. 调 ``build_task_report_docx`` 拿 bytes,用 ``StreamingResponse`` 返回,
     media_type =
     "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

所有样式只许调 docx_style 的助手,不要在这里写样式代码。
"""

from __future__ import annotations

import io
from datetime import datetime

from app.services import docx_style as S

SYSTEM_NAME = "樟巢螟智能检测系统"

_STATUS_CN = {
    "uploading": "上传中",
    "processing": "检测中",
    "completed": "已完成",
    "failed": "失败",
}
_SEVERITY_CN = {"severe": "重度", "medium": "中度", "light": "轻度"}


def build_task_report_docx(*, task, results, nests) -> bytes:
    """生成一份巡检任务 Word 报告,返回 .docx 的二进制内容。

    参数(全部是纯 Python 结构,不是 ORM 对象):
      task: dict —— 键:
        task_name, region_path, area_name, operator, plot_area_mu,
        forestry_sub_compartment, status, created_at(datetime|None),
        completed_at(datetime|None), total_images(int), processed_images(int)
      results: dict|None —— /api/v1/tasks/{id}/results 的返回体
        {"image_stats": {...}, "nest_stats": {...}};任务无结果时传 None
      nests: list[dict] —— 每项键:
        nest_code, longitude, latitude, severity, confidence, detection_count
    """
    from docx import Document

    doc = Document()
    S.setup_document(doc)

    _section_cover(doc, task)
    _section_basic_info(doc, task)
    _section_statistics(doc, results)
    _section_nest_list(doc, nests)

    S.add_page_footer(doc, f"{SYSTEM_NAME} · 巡检报告")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# —— 封面 ————————————————————————————————————————————————
def _section_cover(doc, task):
    S.add_cover(
        doc,
        title=SYSTEM_NAME,
        subtitle="巡检检测报告",
        meta={
            "任务名称": task.get("task_name") or "—",
            "所属区域": task.get("region_path") or "未分配",
            "生成时间": _fmt_dt(datetime.now()),
        },
    )


# —— §1 任务基本信息 ——————————————————————————————————————
def _section_basic_info(doc, task):
    S.add_heading(doc, "一、任务基本信息", level=1)
    rows = [
        ("任务名称", task.get("task_name") or "—"),
        ("所属区域", task.get("region_path") or "未分配"),
        ("巡检区域说明", task.get("area_name") or "—"),
        ("操作员", task.get("operator") or "—"),
        ("地块面积", _fmt_area(task.get("plot_area_mu"))),
        ("林业局小班号", task.get("forestry_sub_compartment") or "—"),
        ("任务状态", _STATUS_CN.get(task.get("status"), task.get("status") or "—")),
        ("创建时间", _fmt_dt(task.get("created_at"))),
        ("完成时间", _fmt_dt(task.get("completed_at"))),
        (
            "影像数量",
            f"{task.get('total_images') or 0} 张"
            f"(已处理 {task.get('processed_images') or 0} 张)",
        ),
    ]
    S.add_kv_table(doc, rows)


# —— §2 检测统计 ————————————————————————————————————————
def _section_statistics(doc, results):
    S.add_heading(doc, "二、检测统计", level=1)
    if not results:
        S.add_body(doc, "该任务暂无检测结果。")
        return

    image_stats = results.get("image_stats") or {}
    S.add_body(doc, "影像处理统计:")
    S.add_data_table(
        doc,
        ["处理图片数", "含香樟树", "含虫巢", "虫巢检测总数"],
        [
            [
                image_stats.get("total_processed", 0),
                image_stats.get("with_camphor_tree", 0),
                image_stats.get("with_nests", 0),
                image_stats.get("total_nest_detections", 0),
            ]
        ],
    )

    nest_stats = results.get("nest_stats") or {}
    S.add_body(doc, "虫巢去重统计:")
    S.add_data_table(
        doc,
        ["去重后虫巢", "重度", "中度", "轻度"],
        [
            [
                nest_stats.get("total_unique", 0),
                nest_stats.get("severe", 0),
                nest_stats.get("medium", 0),
                nest_stats.get("light", 0),
            ]
        ],
    )

    total_unique = nest_stats.get("total_unique", 0) or 0
    severity_rows = []
    for severity, label in (
        ("severe", "重度"),
        ("medium", "中度"),
        ("light", "轻度"),
    ):
        count = nest_stats.get(severity, 0) or 0
        ratio = (count / total_unique) if total_unique else 0
        severity_rows.append(
            [
                label,
                count,
                _fmt_pct(ratio),
                _bar_text(ratio),
            ]
        )

    S.add_body(doc, "严重度分布图表:")
    S.add_data_table(
        doc,
        ["严重度", "数量", "占比", "分布"],
        severity_rows,
    )


# —— §3 虫巢清单 ————————————————————————————————————————
def _section_nest_list(doc, nests):
    S.add_heading(doc, "三、虫巢清单", level=1)
    if not nests:
        S.add_body(doc, "未发现虫巢。")
        return

    rows = []
    for idx, nest in enumerate(nests, start=1):
        rows.append(
            [
                idx,
                nest.get("nest_code") or "—",
                _fmt_coord(nest.get("longitude")),
                _fmt_coord(nest.get("latitude")),
                _SEVERITY_CN.get(
                    nest.get("severity"), nest.get("severity") or "—"
                ),
                _fmt_pct(nest.get("confidence")),
                nest.get("detection_count") or 0,
            ]
        )
    S.add_data_table(
        doc,
        ["编号", "虫巢编码", "经度", "纬度", "严重度", "置信度", "检出次数"],
        rows,
    )


# —— 小工具 ————————————————————————————————————————————————
def _fmt_dt(value):
    """datetime -> 'YYYY-MM-DD HH:MM';空值 -> '—'。"""
    if not value:
        return "—"
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d %H:%M")
    return str(value)


def _fmt_area(value):
    """地块面积 -> 'N 亩';空值 -> '—'。"""
    if value is None:
        return "—"
    return f"{value:g} 亩"


def _fmt_coord(value):
    """经纬度 -> 保留 6 位小数;空值 -> '—'。"""
    if value is None:
        return "—"
    return f"{value:.6f}"


def _fmt_pct(value):
    """置信度 0~1 -> 百分比整数(如 0.87 -> '87%');空值 -> '—'。"""
    if value is None:
        return "—"
    return f"{round(value * 100)}%"


def _bar_text(ratio):
    """用纯文本条形图表达占比,避免在报告中嵌入图片。"""
    filled = round((ratio or 0) * 20)
    return "#" * filled + "-" * (20 - filled)
