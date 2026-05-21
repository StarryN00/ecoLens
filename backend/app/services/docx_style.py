"""python-docx 样式助手 —— T5 巡检报告与 T7 操作手册共用。

设计原则
--------
1. 本模块**只依赖 python-docx 和标准库**,不 import 任何 app 内部模块。
   因此它既能被后端服务(docx_report.py)以 ``app.services.docx_style``
   方式 import,也能被构建脚本(docs/manual/build_manual.py)按文件路径
   直接 import —— 不要在本模块里加任何 ``from app...`` 依赖。
2. 所有生成 Word 的代码都**只允许调用本模块的助手**,不要在别处自己写
   样式 XML —— 这样 T5 报告和 T7 手册的视觉风格才能保持一致。
   要调样式,改本文件顶部的常量即可全局生效。

字体说明
--------
.docx 里存的只是字体「名称」,真正渲染发生在客户端用 Word 打开文档时。
生成方(服务器 / 构建机)无需安装这些中文字体。
"""

from __future__ import annotations

import io
from typing import Iterable, Sequence

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# —— 视觉规范常量(改这里 = 全局改样式)————————————————————————
FONT_HEADING = "微软雅黑"
FONT_BODY = "宋体"

COLOR_BRAND = RGBColor(0x1F, 0x7A, 0x4D)          # 主题绿,与前端 UI 一致
COLOR_TEXT = RGBColor(0x16, 0x25, 0x1D)
COLOR_MUTED = RGBColor(0x60, 0x70, 0x65)
COLOR_TABLE_HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)

HEX_TABLE_HEADER_BG = "1F7A4D"                    # 表头底色(hex,无 #)
HEX_KV_KEY_BG = "EAF3EC"                          # 键值表左列底色

_HEADING_SIZE = {1: 16, 2: 13, 3: 11.5}
_ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}
_CALLOUT_BG = {"tip": "EAF3EC", "warn": "FDF1E7", "info": "EAF1F8"}
_CALLOUT_LABEL = {"tip": "提示", "warn": "注意", "info": "说明"}


# —— 内部工具 ——————————————————————————————————————————————
def _apply_font(run, *, name, size_pt, bold=False, color=None):
    """给一个 run 设置字体。中文要单独设 eastAsia,否则 Word 回落默认宋体。"""
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)


def _shade_cell(cell, hex_color):
    """给表格单元格设背景色。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _write_cell(cell, text, *, bold=False, bg=None, fg=None, align="left"):
    """往单元格写一段文字(会清空原内容)。"""
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = _ALIGN[align]
    run = para.add_run("" if text is None else str(text))
    _apply_font(run, name=FONT_BODY, size_pt=10.5, bold=bold, color=fg or COLOR_TEXT)
    if bg:
        _shade_cell(cell, bg)


def _add_page_number_field(paragraph):
    """在段落里插入一个「当前页码」域(PAGE field)。"""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    _apply_font(run, name=FONT_BODY, size_pt=9, color=COLOR_MUTED)


# —— 对外助手 ——————————————————————————————————————————————
def setup_document(doc):
    """初始化文档:页边距 + 正文默认字体。生成文档第一步就调它。"""
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)
    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(11)
    try:  # 给 Normal 样式补 eastAsia;失败不影响——各助手已逐 run 设字体
        rpr = normal.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), FONT_BODY)
    except Exception:  # noqa: BLE001
        pass


def add_cover(doc, *, title, subtitle="", meta=None):
    """加封面页(末尾自动分页)。meta 是 {标签: 值} 的有序 dict。"""
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_font(
        p.add_run(title), name=FONT_HEADING, size_pt=30, bold=True, color=COLOR_BRAND
    )
    if subtitle:
        sp = doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _apply_font(sp.add_run(subtitle), name=FONT_HEADING, size_pt=16, color=COLOR_TEXT)
    for _ in range(3):
        doc.add_paragraph()
    for key, value in (meta or {}).items():
        mp = doc.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _apply_font(
            mp.add_run(f"{key}：{value}"),
            name=FONT_BODY,
            size_pt=12,
            color=COLOR_TEXT,
        )
    doc.add_page_break()


def add_heading(doc, text, level=1):
    """加章节标题。level: 1=章,2=节,3=小节。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(8)
    _apply_font(
        p.add_run(text),
        name=FONT_HEADING,
        size_pt=_HEADING_SIZE.get(level, 11),
        bold=True,
        color=COLOR_BRAND,
    )
    return p


def add_body(doc, text):
    """加一段正文(1.5 倍行距)。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    _apply_font(p.add_run(text), name=FONT_BODY, size_pt=11, color=COLOR_TEXT)
    return p


def add_bullets(doc, items: Iterable[str]):
    """加无序列表。"""
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        _apply_font(p.add_run(item), name=FONT_BODY, size_pt=11, color=COLOR_TEXT)


def add_callout(doc, text, kind="tip"):
    """加一个提示框(单格底色表格)。kind: tip / warn / info。"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    _shade_cell(cell, _CALLOUT_BG.get(kind, _CALLOUT_BG["info"]))
    cell.text = ""
    para = cell.paragraphs[0]
    _apply_font(
        para.add_run(f"【{_CALLOUT_LABEL.get(kind, '说明')}】"),
        name=FONT_HEADING,
        size_pt=10,
        bold=True,
        color=COLOR_BRAND,
    )
    _apply_font(para.add_run(text), name=FONT_BODY, size_pt=10, color=COLOR_TEXT)


def add_kv_table(doc, rows: Sequence[tuple]):
    """加两列键值表。rows: [(键, 值), ...]。"""
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for key, value in rows:
        cells = table.add_row().cells
        _write_cell(cells[0], key, bold=True, bg=HEX_KV_KEY_BG)
        _write_cell(cells[1], value)
        cells[0].width = Cm(3.6)
        cells[1].width = Cm(11.4)
    return table


def add_data_table(doc, headers: Sequence[str], rows: Sequence[Sequence]):
    """加多列数据表,首行是带底色的表头。"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, head in enumerate(headers):
        _write_cell(
            table.rows[0].cells[i],
            head,
            bold=True,
            bg=HEX_TABLE_HEADER_BG,
            fg=COLOR_TABLE_HEADER_FG,
            align="center",
        )
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            _write_cell(cells[i], value, align="center")
    return table


def add_image_with_caption(doc, image, caption, width_inch=5.6):
    """居中插入一张图 + 下方图注。image 可以是 bytes 或文件路径。"""
    stream = io.BytesIO(image) if isinstance(image, (bytes, bytearray)) else image
    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.add_run().add_picture(stream, width=Inches(width_inch))
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_after = Pt(10)
    _apply_font(cap_p.add_run(caption), name=FONT_BODY, size_pt=9, color=COLOR_MUTED)


def add_page_footer(doc, text):
    """给所有节加页脚:左侧文字 + 居中页码。内容写完后最后调它。"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _apply_font(
            para.add_run(f"{text}    第 "),
            name=FONT_BODY,
            size_pt=9,
            color=COLOR_MUTED,
        )
        _add_page_number_field(para)
        _apply_font(
            para.add_run(" 页"), name=FONT_BODY, size_pt=9, color=COLOR_MUTED
        )
