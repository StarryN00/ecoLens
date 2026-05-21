"""T5 —— 巡检任务 Word 报告导出端点测试。

覆盖:正常导出 / content-type / 附件头 / ownership 隔离 / 不存在 / 鉴权。
共享 fixture(client、auth_headers、second_auth_headers、town_region_id)
见 conftest.py。
"""

import io

from docx import Document

DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def _create_task(client, headers, town_region_id, name="报告测试任务"):
    """建一个挂在 town 级区域上的任务,返回 task_id。"""
    r = client.post(
        "/api/v1/tasks/",
        json={"task_name": name, "region_id": town_region_id},
        headers=headers,
    )
    assert r.status_code == 200, r.text
    return r.json()["id"]


def _docx_text(content: bytes) -> str:
    """把 .docx 里所有段落 + 表格单元格文字拼成一个串,便于断言。"""
    doc = Document(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


class TestTaskReportDocx:
    def test_export_report_ok(self, client, auth_headers, town_region_id):
        name = "陆家嘴巡检报告测试任务"
        task_id = _create_task(client, auth_headers, town_region_id, name)

        r = client.get(
            f"/api/v1/tasks/{task_id}/report.docx", headers=auth_headers
        )
        assert r.status_code == 200, r.text
        assert r.headers["content-type"] == DOCX_MEDIA_TYPE
        assert len(r.content) > 0

        # 内容能被 python-docx 解析,且含关键字段
        text = _docx_text(r.content)
        assert name in text
        assert "任务基本信息" in text
        assert "检测统计" in text
        assert "虫巢清单" in text

    def test_export_report_attachment_header(
        self, client, auth_headers, town_region_id
    ):
        task_id = _create_task(client, auth_headers, town_region_id)
        r = client.get(
            f"/api/v1/tasks/{task_id}/report.docx", headers=auth_headers
        )
        assert r.status_code == 200
        disposition = r.headers.get("content-disposition", "")
        assert "attachment" in disposition
        assert ".docx" in disposition

    def test_export_report_non_owner_404(
        self, client, auth_headers, second_auth_headers, town_region_id
    ):
        """非属主导出他人任务报告 -> 404(与不存在同语义,防枚举)。"""
        task_id = _create_task(client, auth_headers, town_region_id)
        r = client.get(
            f"/api/v1/tasks/{task_id}/report.docx",
            headers=second_auth_headers,
        )
        assert r.status_code == 404

    def test_export_report_admin_can_access(
        self, client, auth_headers, admin_auth_headers, town_region_id
    ):
        """admin 可导出任意任务报告。"""
        task_id = _create_task(client, auth_headers, town_region_id)
        r = client.get(
            f"/api/v1/tasks/{task_id}/report.docx",
            headers=admin_auth_headers,
        )
        assert r.status_code == 200

    def test_export_report_missing_task_404(self, client, auth_headers):
        r = client.get(
            "/api/v1/tasks/no-such-task-id/report.docx", headers=auth_headers
        )
        assert r.status_code == 404

    def test_export_report_requires_auth(
        self, client, auth_headers, town_region_id
    ):
        task_id = _create_task(client, auth_headers, town_region_id)
        r = client.get(f"/api/v1/tasks/{task_id}/report.docx")
        assert r.status_code == 401
