"""T7 —— 操作手册(.docx)构建脚本。

用法:
    python docs/manual/build_manual.py
产物:docs/manual/OPERATION_MANUAL.docx

设计:
- 一章一个函数 ``chapter_NN_xxx(doc)``;``main()`` 顺序调用后保存。
- 章节里的截图放在 docs/manual/assets/,文件名见各章 ``screenshot()`` 调用。
  截图缺失时自动插入占位提示,不影响构建。
- 复用后端 docx_style.py(按文件路径 import,不走 app 包)。
"""

from __future__ import annotations

import os
import sys

HERE = os.path.dirname(os.path.abspath(__file__))
REPO_ROOT = os.path.dirname(os.path.dirname(HERE))
ASSETS_DIR = os.path.join(HERE, "assets")
OUTPUT = os.path.join(HERE, "OPERATION_MANUAL.docx")

sys.path.insert(0, os.path.join(REPO_ROOT, "backend", "app", "services"))

import docx_style as S  # noqa: E402
from docx import Document  # noqa: E402

SYSTEM_NAME = "樟巢螟智能检测系统"


def screenshot(doc, filename, caption):
    """插入截图;文件不存在时插占位提示(补图后自动生效)。"""
    path = os.path.join(ASSETS_DIR, filename)
    if os.path.isfile(path):
        with open(path, "rb") as fh:
            S.add_image_with_caption(doc, fh.read(), caption, width_inch=5.8)
    else:
        S.add_callout(
            doc, f"截图待补:assets/{filename} —— {caption}", kind="warn"
        )


# —— 第 1 章:登录系统 ————————————————————————————————————————
def chapter_01_login(doc):
    S.add_heading(doc, "第 1 章　登录系统", level=1)
    S.add_body(
        doc,
        "樟巢螟智能检测系统是面向林业巡检的无人机影像识别工作台,"
        "用于上传航拍影像、自动识别樟巢螟虫巢并生成巡检报告。"
        "本手册介绍系统各功能的操作步骤,适用于巡检员与系统管理员。",
    )

    S.add_heading(doc, "1.1　打开系统", level=2)
    S.add_body(
        doc,
        "在浏览器中访问系统网址(由管理员提供),进入登录页。"
        "推荐使用 Chrome、Edge 等现代浏览器,以获得最佳显示效果。",
    )
    screenshot(doc, "01_login_page.png", "图 1-1　系统登录页")

    S.add_heading(doc, "1.2　登录步骤", level=2)
    S.add_bullets(
        doc,
        [
            "在「用户名」输入框中填写账号。",
            "在「密码」输入框中填写密码。",
            "点击「登录」按钮;验证通过后即进入巡检任务工作台。",
        ],
    )
    S.add_callout(
        doc,
        "若还没有账号,请联系系统管理员在「用户管理」页面创建;"
        "系统不开放管理员权限的自助注册。",
        kind="tip",
    )

    S.add_heading(doc, "1.3　退出登录", level=2)
    S.add_body(
        doc,
        "点击界面右上角的用户名展开下拉菜单,选择「退出登录」即可安全退出。"
        "在公用电脑上使用完毕后请务必退出登录。",
    )


# —— 第 2 章:角色与权限 ——————————————————————————————————————
def chapter_02_roles(doc):
    S.add_heading(doc, "第 2 章　角色与权限", level=1)
    S.add_body(
        doc,
        "系统分两类角色,登录后左侧导航栏显示的菜单项会随角色不同而变化。",
    )

    S.add_heading(doc, "2.1　普通用户(巡检员)", level=2)
    S.add_bullets(
        doc,
        [
            "新建巡检任务、上传无人机影像。",
            "查看任务的检测结果(虫巢地图、标注影像、虫巢清单)。",
            "导出巡检报告(PDF / Excel / Word)。",
            "修改本人登录密码。",
        ],
    )
    S.add_callout(
        doc,
        "普通用户只能看到本人创建的任务,无法访问他人任务的数据。",
        kind="info",
    )

    S.add_heading(doc, "2.2　系统管理员", level=2)
    S.add_body(doc, "管理员拥有普通用户的全部功能,并额外可以:")
    S.add_bullets(
        doc,
        [
            "进入「区域管理」,维护市 / 区 / 街镇三级行政区域目录。",
            "进入「用户管理」,创建账号、重置密码、启停账号、分配管理员。",
            "查看系统内所有用户的巡检任务。",
        ],
    )


# —— 第 3 章:区域管理(管理员)——————————————————————————————
def chapter_03_region_admin(doc):
    S.add_heading(doc, "第 3 章　区域管理（管理员）", level=1)
    S.add_body(
        doc,
        "系统按「市 → 区 → 街镇」三级行政区域组织巡检任务。"
        "每个巡检任务都必须归属到一个街镇级区域,因此在创建任务之前,"
        "管理员需要先把行政区域目录建好。",
    )
    S.add_body(
        doc,
        "点击左侧导航栏的「区域管理」进入。页面以树形结构展示已有的"
        "市 / 区 / 街镇。",
    )
    screenshot(doc, "03_region_admin.png", "图 3-1　区域管理页")

    S.add_heading(doc, "3.1　新增区域", level=2)
    S.add_bullets(
        doc,
        [
            "新增「市」:点击页面右上角的「新增市」按钮,填写名称后确定。",
            "新增「区」:在某个市节点上选择新增下级,填写区名称。",
            "新增「街镇」:在某个区节点上选择新增下级,填写街镇名称。",
        ],
    )

    S.add_heading(doc, "3.2　重命名与删除", level=2)
    S.add_body(
        doc,
        "在区域节点上可重命名或删除。为保证数据一致,系统对删除有约束:",
    )
    S.add_bullets(
        doc,
        [
            "存在下级区域的节点不能删除,需先删除其所有下级。",
            "已有巡检任务挂载的街镇不能删除。",
        ],
    )
    S.add_callout(
        doc,
        "若已通过初始化脚本导入了完整的市级行政区域目录,通常无需"
        "再手工新增,核对后按需增删即可。",
        kind="tip",
    )


# —— 第 4 章:用户管理(管理员)——————————————————————————————
def chapter_04_user_admin(doc):
    S.add_heading(doc, "第 4 章　用户管理（管理员）", level=1)
    S.add_body(
        doc,
        "点击左侧导航栏的「用户管理」进入。页面以列表展示系统内所有账号,"
        "包含用户名、邮箱、是否管理员、是否启用、创建时间等信息。",
    )
    screenshot(doc, "04_user_admin.png", "图 4-1　用户管理页")

    S.add_heading(doc, "4.1　创建用户", level=2)
    S.add_bullets(
        doc,
        [
            "点击右上角「创建用户」按钮。",
            "在弹窗中填写用户名、初始密码,可选填邮箱。",
            "勾选「管理员」可将该账号直接设为管理员。",
            "确定后,新账号即可用于登录。",
        ],
    )

    S.add_heading(doc, "4.2　管理已有用户", level=2)
    S.add_bullets(
        doc,
        [
            "重置密码:为忘记密码的用户设置新密码。",
            "启用 / 停用:停用的账号将无法登录。",
            "设为 / 取消管理员:调整账号的管理员权限。",
        ],
    )
    S.add_callout(
        doc,
        "请为管理员账号设置足够强度的密码,并定期更换,避免使用 "
        "admin123 一类弱口令。",
        kind="warn",
    )


# —— 第 5 章:新建巡检任务 ————————————————————————————————————
def chapter_05_create_task(doc):
    S.add_heading(doc, "第 5 章　新建巡检任务", level=1)
    S.add_body(
        doc,
        "点击左侧导航栏的「新建任务」进入。创建流程分三步:"
        "基本信息 → 上传图片 → 完成。",
    )
    screenshot(doc, "05_create_task.png", "图 5-1　新建巡检任务 - 基本信息")

    S.add_heading(doc, "5.1　填写基本信息", level=2)
    S.add_bullets(
        doc,
        [
            "任务名称:必填,建议含地点与日期,便于检索。",
            "所属区域:必填,通过三级级联选择器依次选「市 / 区 / 街镇」,"
            "必须选到街镇级。",
            "巡检区域说明:选填,对巡检范围的文字补充描述。",
            "操作员:选填,本次巡检的操作人员姓名。",
            "地块面积:选填,单位为「亩」。",
            "林业局小班号:选填,如 A-12-3。",
        ],
    )
    S.add_callout(
        doc,
        "若区域级联选择器中没有可选项,说明尚未建立行政区域目录,"
        "请联系管理员在「区域管理」中创建(见第 3 章)。",
        kind="tip",
    )

    S.add_heading(doc, "5.2　上传无人机影像", level=2)
    S.add_body(
        doc,
        "进入「上传图片」步骤,选择本次巡检的无人机航拍影像批量上传。"
        "上传完成后,系统会自动对影像进行 AI 检测,无需手动触发。",
    )
    screenshot(doc, "05_upload.png", "图 5-2　上传无人机影像")

    S.add_heading(doc, "5.3　完成", level=2)
    S.add_body(
        doc,
        "上传完成后任务即创建成功,可进入任务详情页查看检测进度与结果。",
    )


# —— 第 6 章:查看检测结果 ————————————————————————————————————
def chapter_06_view_results(doc):
    S.add_heading(doc, "第 6 章　查看检测结果", level=1)
    S.add_body(
        doc,
        "在「巡检任务」列表中点击某个任务即进入任务详情页。页面顶部是"
        "图片总数、推理进度、去重虫巢、重度风险等关键指标卡,下方是四个"
        "标签页。",
    )

    S.add_heading(doc, "6.1　概览", level=2)
    S.add_body(
        doc,
        "展示任务档案(区域、地块、操作员等)与检测统计汇总,"
        "底部提供「导出报告」入口。",
    )
    screenshot(doc, "06_overview.png", "图 6-1　任务详情 - 概览")

    S.add_heading(doc, "6.2　地图", level=2)
    S.add_body(doc, "在地图上按 GPS 坐标展示各虫巢的空间分布。")
    screenshot(doc, "06_map.png", "图 6-2　任务详情 - 虫巢分布地图")

    S.add_heading(doc, "6.3　图片", level=2)
    S.add_body(
        doc,
        "列出任务的每张影像及其检测结果。点击影像可打开标注查看器,"
        "查看在原图上标注的虫巢检测框。",
    )
    S.add_callout(
        doc,
        "虫巢位置以红色加粗方框标注。查看器默认加载压缩图以加快速度,"
        "点击「查看原图」可加载全分辨率影像。",
        kind="info",
    )
    screenshot(doc, "06_images.png", "图 6-3　任务详情 - 图片标注")

    S.add_heading(doc, "6.4　虫巢", level=2)
    S.add_body(
        doc,
        "展示去重后的虫巢清单,含编号、经纬度、严重度、置信度等。"
        "同一虫巢在多张影像中重复出现时会自动合并去重。",
    )
    screenshot(doc, "06_nests.png", "图 6-4　任务详情 - 虫巢清单")


# —— 第 7 章:导出报告 ——————————————————————————————————————
def chapter_07_export_report(doc):
    S.add_heading(doc, "第 7 章　导出报告", level=1)
    S.add_body(
        doc,
        "在任务详情页「概览」标签底部点击「导出报告」按钮,弹出格式选择窗口。",
    )
    screenshot(doc, "07_export_modal.png", "图 7-1　导出报告窗口")

    S.add_heading(doc, "7.1　报告格式", level=2)
    S.add_bullets(
        doc,
        [
            "PDF 报告:含基本信息、统计与虫巢列表,适合打印归档。",
            "Excel 报告:基本信息与虫巢列表分表,便于二次数据处理。",
            "Word 报告:完整巡检报告,含任务信息、检测统计、虫巢清单"
            "以及标注影像附录,由服务端生成,适合正式交付。",
            "CSV 数据:导出虫巢点位的纯数据表。",
        ],
    )
    S.add_callout(
        doc,
        "点击对应按钮后浏览器会自动下载文件;Word 报告生成稍慢,"
        "请耐心等待下载完成。",
        kind="tip",
    )


# —— 第 8 章:修改密码 ——————————————————————————————————————
def chapter_08_change_password(doc):
    S.add_heading(doc, "第 8 章　修改密码", level=1)
    S.add_body(
        doc,
        "任何用户都可以修改本人的登录密码。",
    )
    S.add_bullets(
        doc,
        [
            "点击界面右上角的用户名,展开下拉菜单。",
            "选择「修改密码」。",
            "在弹窗中输入当前密码与新密码,确定后即生效。",
        ],
    )
    screenshot(doc, "08_change_password.png", "图 8-1　修改密码窗口")
    S.add_callout(
        doc,
        "若忘记当前密码无法自助修改,请联系管理员在「用户管理」中重置。",
        kind="tip",
    )


# —— 第 9 章:常见问题与故障排查 ——————————————————————————————
def chapter_09_faq(doc):
    S.add_heading(doc, "第 9 章　常见问题与故障排查", level=1)

    S.add_heading(doc, "登录失败", level=2)
    S.add_body(
        doc,
        "请核对用户名与密码大小写;若账号被停用或不存在,请联系管理员。",
    )

    S.add_heading(doc, "无法新建任务 / 区域选不到", level=2)
    S.add_body(
        doc,
        "创建任务必须选完整的三级区域并选到街镇级。若级联选择器为空,"
        "说明尚未建立行政区域目录,请联系管理员在「区域管理」中创建。",
    )

    S.add_heading(doc, "检测结果为空 / 虫巢数为 0", level=2)
    S.add_body(
        doc,
        "可能影像中确实没有虫巢;若怀疑是 AI 模型未就绪,请联系管理员"
        "确认服务器端的检测模型已正确部署。",
    )

    S.add_heading(doc, "图片加载慢", level=2)
    S.add_body(
        doc,
        "系统默认加载压缩后的影像以加快显示。确需查看全分辨率原图时,"
        "在标注查看器中点击「查看原图」。",
    )

    S.add_callout(
        doc,
        "遇到本章未涵盖的问题,请联系系统管理员或技术支持。",
        kind="info",
    )


def main():
    doc = Document()
    S.setup_document(doc)
    S.add_cover(
        doc,
        title=SYSTEM_NAME,
        subtitle="操作手册",
        meta={"文档版本": "v1.0", "适用对象": "巡检员 / 系统管理员"},
    )

    chapter_01_login(doc)
    chapter_02_roles(doc)
    chapter_03_region_admin(doc)
    chapter_04_user_admin(doc)
    chapter_05_create_task(doc)
    chapter_06_view_results(doc)
    chapter_07_export_report(doc)
    chapter_08_change_password(doc)
    chapter_09_faq(doc)

    S.add_page_footer(doc, f"{SYSTEM_NAME} · 操作手册")

    doc.save(OUTPUT)
    print(f"OK: 已生成 {OUTPUT}")


if __name__ == "__main__":
    main()
